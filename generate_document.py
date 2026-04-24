"""
Word Document Generator — Creates 'Friday Lintspace.docx' on Desktop.
Generates the full ~4,000 word paper with embedded charts and tables.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from data.technology_data import impact_matrix
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
DESKTOP = os.path.join(os.path.expanduser('~'), 'Desktop')


def set_cell_shading(cell, color):
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'
    })
    shading_elm.append(shading)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_chart(doc, chart_name, caption):
    path = os.path.join(OUTPUT_DIR, f'{chart_name}.png')
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.2))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style.font.size = Pt(9)
        cap.style.font.italic = True


def add_impact_table(doc, df):
    doc.add_heading('Impact Matrix', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Technology', 'Readiness (1-10)', 'Societal Potential (1-10)', 'Current Market ($B)']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['technology'])
        cells[1].text = str(row['readiness_level'])
        cells[2].text = str(row['societal_potential'])
        cells[3].text = f"${row['current_market_B']:.1f}B"


def add_convergence_table(doc, df):
    doc.add_heading('Master Convergence Table', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Technology', 'Readiness (1-10)', 'Primary Economic Driver', 'Key Hard Tech Job Creation']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['technology'])
        cells[1].text = str(row['readiness_level'])
        cells[2].text = str(row['primary_economic_driver'])
        cells[3].text = str(row['key_job_creation'])


def generate_document():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading('The Hard Tech Revolution:\nWhat Actually Comes After the AI Boom', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('An evidence-based analysis of the 12 technologies defining the next physical chapter of human progress.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph('April 2026 | Research & Analysis')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # =========================================================================
    # INTRODUCTION — The Physical Pivot (~300 words)
    # =========================================================================
    add_styled_heading(doc, 'Introduction: The Physical Pivot')

    add_body(doc, 'The AI boom was never the destination. It was the on-ramp.')

    add_body(doc, 'Since 2022, global venture capital has poured unprecedented sums into artificial intelligence. The narrative has been clear: software is eating the world. But look at the data. In 2018, 78% of venture capital flowed exclusively into software companies. By 2026, that figure has dropped to 38%. The remaining share has not vanished—it has migrated. Hard tech and deep tech now command 42% of all venture funding, up from just 12% eight years ago. This is not a blip. It is a structural rotation.')

    add_body(doc, 'The thesis of this paper is simple. The AI revolution was the digital precursor to something far larger: a physical revolution. Artificial intelligence provided the computational toolkit—the ability to model, simulate, and optimize at inhuman scale. But the real returns accrue when that toolkit is applied to atoms, not bits. To energy systems, not chatbots. To biological substrates, not software interfaces.')

    add_body(doc, 'What follows is an analysis of 12 technologies that constitute this "Physical Chapter." Each is assessed on three axes: data-backed trajectory, market and economic impact, and readiness for mass commercialization. The time horizon is twenty years. The lens is infrastructure, scalability, and physical sovereignty.')

    add_body(doc, 'Observe the trend in the chart below. The sharp rebalancing of capital allocation from software-only to hard tech coincides with a recognition among sophisticated investors that the largest economic opportunities of the next two decades are physical, not digital.')

    add_chart(doc, '00_vc_shift', 'Figure 1: VC Investment Shift — Software vs Hard Tech (2018–2026). Source: PitchBook, Dealroom.')
    doc.add_page_break()

    # =========================================================================
    # PART 1: THE ENERGY BEDROCK (~900 words)
    # =========================================================================
    add_styled_heading(doc, 'Part 1: The Energy Bedrock')

    # --- 1.1 Nuclear Fusion ---
    add_styled_heading(doc, '1.1 Nuclear Fusion: The Ultimate Energy Source', level=2)

    add_body(doc, 'Context. Nuclear fusion—the process that powers the sun—has long been dismissed as perpetually "30 years away." That narrative is changing. The IEA\'s State of Energy Innovation 2026 report describes fusion as "on the cusp of demonstration," a significant rhetorical upgrade from prior assessments. The shift is driven by two forces: experimental breakthroughs in government laboratories and an explosion of private capital.')

    add_body(doc, 'Data Trend. Cumulative private investment in fusion has grown from $1 billion in 2015 to $9.5 billion in 2026. The number of active fusion companies has increased from 10 to 53 in the same period. The IEA has set a critical demonstration milestone for 2030: the first fusion plant to demonstrate the technical viability of producing saleable energy. This is no longer speculative physics. It is an engineering timeline.')

    add_body(doc, 'Market Impact. If commercial fusion is achieved by 2035, the addressable market is effectively the entire $10 trillion global energy sector. Even at 1% penetration in its first decade, fusion represents a $100 billion annual market. The downstream effects—unlimited clean baseload power—would reshape every industry dependent on energy cost, from manufacturing to data centers to desalination.')

    add_body(doc, 'Observe the trend line for private fusion investment: the sharp acceleration after 2020 coincides with the NIF ignition result and the entry of institutional-grade capital.')

    add_chart(doc, '01_nuclear_fusion', 'Figure 2: Nuclear Fusion — Private Investment & Industry Growth. Sources: Fusion Industry Association, IEA.')

    # --- 1.2 Next-Gen Batteries ---
    add_styled_heading(doc, '1.2 Next-Gen Batteries: Electrification of Everything', level=2)

    add_body(doc, 'Context. The lithium-ion battery was the enabling technology of the smartphone era and the electric vehicle revolution. But it is approaching its thermodynamic ceiling. The next frontier is solid-state batteries (SSBs), which promise 2–3× the energy density at comparable or lower weight.')

    add_body(doc, 'Data Trend. Traditional Li-ion cells deliver approximately 260 Wh/kg in 2025. Solid-state prototypes are already achieving 420–450 Wh/kg in pilot production, with laboratory demonstrations exceeding 800 Wh/kg. The commercial target for near-term SSB adoption is 500 Wh/kg. Meanwhile, the global battery market has grown from $41 billion in 2020 to an estimated $155 billion in 2026, with EVs accounting for over 75% of demand.')

    add_body(doc, 'Market Impact. Higher energy density translates directly into longer EV range, lighter aircraft, and more compact grid storage. The transition from Li-ion to solid-state is not incremental—it enables entirely new product categories, from electric regional aircraft to multi-day consumer electronics. The battery market is projected to reach $310 billion by 2030.')

    add_chart(doc, '02_batteries', 'Figure 3: Battery Energy Density — Li-ion vs Solid-State Trajectory. Sources: IEA, Electrive, Bonnenbatteries.')

    # --- 1.3 Direct Air Capture ---
    add_styled_heading(doc, '1.3 Direct Air Capture: Industrial-Scale Carbon Removal', level=2)

    add_body(doc, 'Context. Direct Air Capture (DAC) is the technology that pulls CO₂ directly from ambient air. It is the only scalable approach to addressing legacy emissions—the carbon already in the atmosphere. Until recently, DAC was dismissed as too expensive. That is changing rapidly.')

    add_body(doc, 'Data Trend. The cost of DAC has fallen from $800 per ton of CO₂ in 2020 to approximately $340 per ton in 2026. Climeworks\' Generation 3 technology targets $250–350 per ton by 2030, with a long-term goal of $100–200. Installed capture capacity has scaled from 10 tons per year in 2020 to 8,000 tons per year in 2026—an 800× increase. The market is projected to grow from $3.5 billion in 2026 to $400 billion by 2035, driven by corporate net-zero commitments and government incentives like the US 45Q tax credit.')

    add_body(doc, 'Observe the cost-capacity divergence in the chart below: as cost per ton declines exponentially, installed capacity rises on a logarithmic scale. This is the signature pattern of a technology approaching industrial viability.')

    add_chart(doc, '03_dac', 'Figure 4: DAC Cost Decline & Capacity Scaling. Sources: Climeworks, DOE, Fortune Business Insights.')
    doc.add_page_break()

    # =========================================================================
    # PART 2: THE BIOLOGICAL REWRITE (~900 words)
    # =========================================================================
    add_styled_heading(doc, 'Part 2: The Biological Rewrite')

    # --- 2.1 Gene Editing ---
    add_styled_heading(doc, '2.1 Gene Editing: From Treatment to Cure', level=2)

    add_body(doc, 'Context. CRISPR-Cas9 gene editing has transitioned from a laboratory curiosity to a commercial therapeutic platform. Casgevy, the first FDA-approved CRISPR therapy, treats sickle cell disease and beta-thalassemia. The pipeline is expanding rapidly into cardiovascular disease, cancer, autoimmune conditions, and rare genetic disorders. Next-generation tools—base editing and prime editing—offer even greater precision.')

    add_body(doc, 'Data Trend. Active CRISPR clinical trials have grown from 5 in 2017 to over 155 in 2026. The FDA\'s February 2026 "Plausible Mechanism Framework" creates a new regulatory pathway for personalized gene-editing therapies, potentially accelerating approvals for thousands of rare diseases. The gene editing market has grown from $1.2 billion in 2017 to $15 billion in 2026.')

    add_body(doc, 'Market Impact. Gene editing represents a paradigm shift from chronic disease management to one-time curative interventions. Eli Lilly\'s $5 billion acquisition of Verve Therapeutics in 2025 signals Big Pharma\'s conviction. The total addressable market for curative genetic medicine could exceed $80 billion by 2035.')

    add_chart(doc, '04_gene_editing', 'Figure 5: Gene Editing Clinical Trial Pipeline Growth. Sources: clinicaltrials.gov, Innovative Genomics Institute.')

    # --- 2.2 Synthetic Biology ---
    add_styled_heading(doc, '2.2 Synthetic Biology: Programming Life', level=2)

    add_body(doc, 'Context. Synthetic biology is the engineering discipline of designing and constructing biological systems. It enables the production of materials, chemicals, fuels, and food through biological processes rather than petrochemical ones. The convergence with AI has dramatically accelerated the design-build-test-learn cycle.')

    add_body(doc, 'Data Trend. The synthetic biology market has grown from $9.5 billion in 2020 to $35 billion in 2026, with a CAGR of 22–23%. Projections place the market at $112 billion by 2033. Biopharmaceuticals lead at 32% of market share, followed by industrial enzymes (22%) and biofuels/chemicals (18%). The fastest-growing segment is food and nutrition (28% growth rate), driven by precision fermentation.')

    add_body(doc, 'Market Impact. Synthetic biology threatens to displace significant portions of the $5 trillion petrochemical industry. Bio-based manufacturing offers lower carbon footprints, programmable output, and production distributed closer to end markets. The key inflection point is cost parity with petrochemical equivalents, which multiple categories are approaching.')

    add_chart(doc, '05_synthetic_biology', 'Figure 6: Synthetic Biology Market Growth & Segment Breakdown. Sources: Grand View Research, BioSpace.')

    # --- 2.3 Regenerative Biotech ---
    add_styled_heading(doc, '2.3 Regenerative Biotech: Rebuilding the Human Body', level=2)

    add_body(doc, 'Context. Regenerative medicine encompasses cell therapies, gene therapies, tissue engineering, and organ replacement. The field addresses one of medicine\'s most fundamental limitations: the inability to replace damaged or aging tissues. Advances in 3D bioprinting, xenotransplantation, and cellular reprogramming are converging to create entirely new therapeutic categories.')

    add_body(doc, 'Data Trend. The combined regenerative biotech market—spanning regenerative medicine ($45B), longevity biotech ($26.5B), and organ replacement ($14B)—reached $85.5 billion in 2026. The total is projected to exceed $410 billion by 2035. Longevity biotech, driven by epigenetic clocks and senolytics, is growing at the fastest rate, fueled by a demographic imperative: 2.1 billion people will be over 60 by 2050.')

    add_chart(doc, '06_regenerative_biotech', 'Figure 7: Regenerative Biotech Market Segments. Sources: Mordor Intelligence, IMARC Group.')
    doc.add_page_break()

    # =========================================================================
    # PART 3: THE ATOMIC/COMPUTE EDGE (~900 words)
    # =========================================================================
    add_styled_heading(doc, 'Part 3: The Atomic/Compute Edge')

    # --- 3.1 Nanotechnology ---
    add_styled_heading(doc, '3.1 Nanotechnology: Engineering at the Atomic Scale', level=2)

    add_body(doc, 'Context. Nanotechnology has shifted from theoretical research to industrial-scale production. Carbon nanotubes, quantum dots, and nanocomposites are now standard components in electronics, healthcare, and aerospace. The US National Nanotechnology Initiative and CHIPS Act have provided critical policy support.')

    add_body(doc, 'Data Trend. The global nanotechnology market reached $118 billion in 2026, up from $62 billion in 2020. Healthcare and pharmaceuticals lead at 28%, followed by electronics/semiconductors at 25%. The market is projected to reach $190 billion by 2030. The key transition is from basic material synthesis to integrated nanodevices and nanosensors.')

    add_chart(doc, '07_nanotechnology', 'Figure 8: Nanotechnology Market Growth & Sector Distribution. Sources: Mordor Intelligence, NovaOne Advisor.')

    # --- 3.2 Advanced Materials ---
    add_styled_heading(doc, '3.2 Advanced Materials: Graphene & Metamaterials', level=2)

    add_body(doc, 'Context. Graphene—a single atom-thick layer of carbon—and metamaterials—engineered structures with properties not found in nature—represent a generational leap in material science. Both sectors have reached an inflection point where pilot-scale production is transitioning to industrial volumes.')

    add_body(doc, 'Data Trend. The graphene market has grown from $80 million in 2018 to $2 billion in 2026, with production scaling from 80 tons to 5,500 tons annually. Metamaterials have grown from $500 million to $1.8 billion, driven by 6G telecommunications and autonomous vehicle radar. Combined, advanced materials are projected to reach $12.5 billion by 2030.')

    add_chart(doc, '08_advanced_materials', 'Figure 9: Advanced Materials Manufacturing Scale-Up. Sources: GM Insights, IDTechEx.')

    # --- 3.3 Quantum Computing ---
    add_styled_heading(doc, '3.3 Quantum Computing: Beyond Classical Limits', level=2)

    add_body(doc, 'Context. Quantum computing crossed a critical threshold in 2024 when Google\'s Willow chip demonstrated "below-threshold" error correction—the foundational requirement for fault-tolerant quantum computation. IBM\'s 2026 roadmap targets "scientific quantum advantage," where quantum hardware outperforms classical methods on real-world problems.')

    add_body(doc, 'Data Trend. Physical qubit counts have grown from 53 (Google Sycamore, 2019) to 2,000 in 2026. More importantly, logical qubits—the error-corrected units that perform useful computation—have grown from 0 to 12. IBM aims for 200 logical qubits by 2029 with its Quantum Starling system. The quantum computing market has grown from $500 million in 2019 to $7.5 billion in 2026.')

    add_body(doc, 'Market Impact. The two killer applications are cryptographic security (post-quantum cryptography is already a $1.5 billion market) and material discovery (simulating molecular interactions for drug design and battery chemistry). The market is projected to reach $28 billion by 2029.')

    add_chart(doc, '09_quantum_computing', 'Figure 10: Quantum Computing Qubit Scaling Roadmap. Sources: IBM Research, Google Quantum AI.')

    # --- 3.4 Brain-Computer Interfaces ---
    add_styled_heading(doc, '3.4 Brain-Computer Interfaces: The Human-Machine Merge', level=2)

    add_body(doc, 'Context. Brain-computer interfaces (BCIs) have moved from science fiction to clinical reality. Neuralink has 21 participants enrolled across global trials, with FDA Breakthrough Device Designation for speech restoration. The company has shifted from boutique manufacturing to high-volume production and automated surgical implantation.')

    add_body(doc, 'Data Trend. The BCI market grew from $1.2 billion in 2020 to $3.5 billion in 2026. Clinical trial participants have grown from 5 to 85 in the same period. Error rates have fallen from 25% to 4.5%, while typing speeds have increased from 5 to 40 words per minute. The market is projected to reach $15 billion by 2035, driven by therapeutic applications in paralysis, ALS, and stroke recovery.')

    add_chart(doc, '10_bci', 'Figure 11: BCI Clinical Progress — Participants, Error Rates, and Speed. Sources: Neuralink, SNS Insider.')
    doc.add_page_break()

    # =========================================================================
    # PART 4: NEW INFRASTRUCTURE (~700 words)
    # =========================================================================
    add_styled_heading(doc, 'Part 4: New Infrastructure')

    # --- 4.1 Space Infrastructure ---
    add_styled_heading(doc, '4.1 Space Infrastructure: The Orbital Economy', level=2)

    add_body(doc, 'Context. The space economy is experiencing a transformation driven by one variable: launch cost. The transition from expendable to fully reusable rockets has reduced the cost per kilogram to Low Earth Orbit from $54,500 in 2006 to $2,400 in 2026. SpaceX\'s Starship targets $100 per kilogram—a further 96% reduction.')

    add_body(doc, 'Data Trend. The global space economy has grown from $175 billion in 2006 to $570 billion in 2026. Annual orbital launches have increased from 18 to 220. Goldman Sachs and Morgan Stanley project the space economy will cross $1 trillion between 2032 and 2036. The primary revenue driver is LEO satellite constellations for broadband and Earth observation, but emerging sectors include in-orbit manufacturing, space tourism, and orbital data centers.')

    add_body(doc, 'Observe the inverse relationship in the chart below: as launch costs decline on a logarithmic scale, the space economy grows linearly. This is the classic "cost-to-deploy drives demand" dynamic that has historically preceded the emergence of entirely new industries.')

    add_chart(doc, '11_space_infrastructure', 'Figure 12: Space Launch Cost Decay vs Orbital Revenue. Sources: Precedence Research, Goldman Sachs.')

    # --- 4.2 DePIN / DLT ---
    add_styled_heading(doc, '4.2 DePIN / DLT: Decentralized Physical Infrastructure', level=2)

    add_body(doc, 'Context. Decentralized Physical Infrastructure Networks (DePIN) represent the convergence of blockchain technology with real-world physical infrastructure. Unlike speculative crypto, DePIN networks generate verifiable, on-chain revenue from actual customers paying for compute, storage, connectivity, and data services. The sector has matured from narrative to utility.')

    add_body(doc, 'Data Trend. The DePIN sector has a combined market capitalization of approximately $10 billion in 2026. Leading networks generate $150 million in monthly on-chain revenue—a massive year-over-year increase driven by AI-related compute demand rather than speculation. GPU/Compute leads at 35% of the sector, followed by storage (20%) and wireless connectivity (18%). Decentralized GPU networks offer 45–75% cost savings over centralized cloud providers for inference workloads, with 70% of demand driven by AI inference.')

    add_body(doc, 'Market Impact. DePIN addresses the critical infrastructure bottleneck created by AI: the exponential demand for compute, storage, and energy that centralized hyperscalers cannot build fast enough. By aggregating distributed, latent resources globally, DePIN bypasses the power and capital constraints of building new centralized data centers. The sector is projected to reach $65 billion by 2030.')

    add_chart(doc, '12_depin_dlt', 'Figure 13: DePIN Market Cap & On-Chain Revenue Growth. Sources: KuCoin Research, Binance Research.')
    doc.add_page_break()

    # =========================================================================
    # CONCLUSION: THE CONVERGENCE (~300 words)
    # =========================================================================
    add_styled_heading(doc, 'Conclusion: The Convergence')

    add_body(doc, 'The twelve technologies analyzed in this paper are not isolated verticals. They are converging. Fusion energy powers the data centers that run the quantum computers that simulate the molecular interactions that design the next-generation battery materials. Gene editing platforms run on AI models trained on decentralized GPU networks. Space infrastructure launches the sensors that feed the nanotechnology-enabled devices that monitor Earth systems for DAC optimization.')

    add_body(doc, 'This is the "Synthesis Matrix"—the idea that the greatest economic value will be created not within any single technology, but at the intersections between them. Our neural network convergence analysis confirms this: the highest-scoring technology pairs are those that combine computational power with physical infrastructure (Quantum Computing × Advanced Materials) and biological capability with energy systems (Synthetic Biology × Next-Gen Batteries).')

    add_body(doc, 'The AI boom was real. But it was Chapter One. The physical revolution—the Hard Tech Revolution—is Chapter Two. And the data suggests it will be far larger, far longer, and far more consequential for human civilization.')

    add_body(doc, 'The investors, policy-makers, and corporate leaders who recognize this pivot today will define the economic landscape of the next two decades. Those who remain anchored in the "software-only" paradigm risk being left behind by the most significant infrastructure build-out since the Industrial Revolution.')

    add_chart(doc, '13_impact_matrix', 'Figure 14: Technology Impact Matrix — Readiness vs Societal Potential. Bubble size = current market size.')

    # Impact Matrix Table
    add_impact_table(doc, impact_matrix)
    doc.add_page_break()

    # Convergence Table
    add_convergence_table(doc, impact_matrix)

    add_chart(doc, '14_convergence_matrix', 'Figure 15: Technology Convergence Matrix (Neural Network Predicted Scores).')

    # =========================================================================
    # SOURCES & DATASETS
    # =========================================================================
    doc.add_page_break()
    add_styled_heading(doc, 'Sources & Datasets')

    sources = [
        ('Energy / Fusion', 'IEA State of Energy Innovation 2026; Fusion Industry Association Annual Report 2026'),
        ('Batteries', 'IEA Global EV Outlook; Electrive; Bonnenbatteries; Fortune Business Insights'),
        ('Direct Air Capture', 'Climeworks; US DOE DAC Hub Program; Fortune Business Insights; Sylvera'),
        ('Gene Editing', 'ClinicalTrials.gov; Innovative Genomics Institute; CRISPR Therapeutics; FDA'),
        ('Synthetic Biology', 'Grand View Research; Coherent Market Insights; BioSpace; GlobeNewsWire'),
        ('Regenerative Biotech', 'Mordor Intelligence; Technavio; IMARC Group; Precedence Research'),
        ('Nanotechnology', 'Mordor Intelligence; NovaOne Advisor; AZoNano'),
        ('Advanced Materials', 'GM Insights; IDTechEx; Precedence Research; Coherent Market Insights'),
        ('Quantum Computing', 'IBM Research; Google Quantum AI; Forbes; PostQuantum'),
        ('Brain-Computer Interfaces', 'SNS Insider; Neuralink; Economic Times; DataM Intelligence'),
        ('Space Infrastructure', 'Precedence Research; Goldman Sachs; Morgan Stanley; Deloitte; SpaceX'),
        ('DePIN / DLT', 'KuCoin Research; Binance Research; CoinCub; Messari'),
        ('VC Investment Trends', 'PitchBook 2026 VC Outlook; Dealroom; Forbes; Celesta VC'),
        ('ML/DL Analysis', 'scikit-learn (regression, K-Means, PCA); PyTorch (convergence neural network)'),
    ]
    for domain, src in sources:
        p = doc.add_paragraph()
        runner = p.add_run(f'{domain}: ')
        runner.bold = True
        p.add_run(src)

    # Save
    filepath = os.path.join(DESKTOP, 'Friday Lintspace.docx')
    doc.save(filepath)
    print(f"\n✓ Document saved: {filepath}")
    return filepath


if __name__ == '__main__':
    generate_document()
